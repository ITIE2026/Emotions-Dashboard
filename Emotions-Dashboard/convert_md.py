import re, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

md_path = r'C:\Users\asus\Documents\dashboard\Emotions-Dashboard\MANAGER_EEG_DASHBOARD_TECHNICAL_REVIEW.md'
out_path = r'C:\Users\asus\Documents\dashboard\Emotions-Dashboard\MANAGER_EEG_DASHBOARD_TECHNICAL_REVIEW.docx'

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')
    
    # Heading 1
    if line.startswith('# ') and not line.startswith('## '):
        p = doc.add_heading(line[2:], level=1)
    # Heading 2
    elif line.startswith('## ') and not line.startswith('### '):
        p = doc.add_heading(line[3:], level=2)
    # Heading 3
    elif line.startswith('### ') and not line.startswith('#### '):
        p = doc.add_heading(line[4:], level=3)
    # Heading 4
    elif line.startswith('#### '):
        p = doc.add_heading(line[5:], level=4)
    # Table row  
    elif line.startswith('|'):
        # skip separator rows
        if re.match(r'\|[-| ]+\|', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        # Check if this is first row (check next line for separator)
        is_header = i+1 < len(lines) and re.match(r'\|[-| ]+\|', lines[i+1])
        if not hasattr(doc, '_current_table') or doc._current_table is None:
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Table Grid'
            doc._current_table = table
            doc._table_col_count = len(cells)
            row = table.rows[0]
        else:
            row = doc._current_table.add_row()
        for j, cell_text in enumerate(cells[:doc._table_col_count]):
            cell = row.cells[j]
            cell.text = cell_text
            if is_header:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        i += 1
        continue
    # Bullet list
    elif line.startswith('- ') or line.startswith('* '):
        doc._current_table = None
        p = doc.add_paragraph(line[2:], style='List Bullet')
    # Numbered list
    elif re.match(r'^\d+\. ', line):
        doc._current_table = None
        p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
    # Code block
    elif line.startswith('```'):
        doc._current_table = None
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].rstrip('\n').startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        p = doc.add_paragraph('\n'.join(code_lines))
        p.style = doc.styles['No Spacing']
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    # Empty line
    elif line.strip() == '':
        doc._current_table = None
        doc.add_paragraph('')
    # Normal paragraph
    else:
        doc._current_table = None
        # Strip inline backticks for plain text
        text = re.sub(r'`([^`]+)`', r'\1', line)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        p = doc.add_paragraph(text)
    
    i += 1

doc.save(out_path)
print('Saved to', out_path)
